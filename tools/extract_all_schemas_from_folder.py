import re
import yaml
import argparse
from pathlib import Path
from docx import Document

def extract_merge_fields(docx_path):
    field_pattern = re.compile(r"<<\s*(.*?)\s*>>")
    fields = set()

    def extract_from_text(text):
        matches = field_pattern.findall(text)
        for match in matches:
            fields.add(match.strip())

    doc = Document(docx_path)
    for para in doc.paragraphs:
        extract_from_text(para.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                extract_from_text(cell.text)

    return sorted(fields)

def generate_schema(fields, template_name):
    return {
        "schema": template_name,
        "required_fields": fields
    }

def write_schema(schema, output_path):
    with open(output_path, "w", encoding="utf-8") as f:
        yaml.dump(schema, f, sort_keys=False)
    print(f"✅ Saved schema: {output_path}")

def main():
    parser = argparse.ArgumentParser(description="Extract schemas from all DOCX templates in a folder")
    parser.add_argument("folder", help="Path to folder containing .docx templates")
    parser.add_argument("--output", help="Folder to save schema files", default="schemas")
    args = parser.parse_args()

    input_folder = Path(args.folder)
    output_folder = Path(args.output)
    output_folder.mkdir(parents=True, exist_ok=True)

    for docx_file in input_folder.glob("*.docx"):
        fields = extract_merge_fields(docx_file)
        schema_name = docx_file.stem
        schema_data = generate_schema(fields, schema_name)
        output_path = output_folder / f"{schema_name}.schema.yml"
        write_schema(schema_data, output_path)

if __name__ == "__main__":
    main()