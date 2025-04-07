import re
import argparse
from pathlib import Path
from docx import Document
import yaml

def extract_merge_fields(docx_path):
    doc = Document(docx_path)
    field_pattern = re.compile(r"<<\s*(.*?)\s*>>")
    fields = set()

    def extract_from_text(text):
        matches = field_pattern.findall(text)
        for match in matches:
            fields.add(match.strip())

    for para in doc.paragraphs:
        extract_from_text(para.text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                extract_from_text(cell.text)

    return sorted(fields)

def generate_schema(fields, template_name):
    return {
        "schema": template_name,
        "required_fields": fields
    }

def save_schema(schema, output_path):
    with open(output_path, "w", encoding="utf-8") as f:
        yaml.dump(schema, f, sort_keys=False)
    print(f"✅ Schema saved to: {output_path}")

def main():
    parser = argparse.ArgumentParser(description="Extract merge fields from DOCX and create schema")
    parser.add_argument("input", help="Path to .docx file")
    parser.add_argument("-o", "--output", help="Schema output .yml path (optional)")
    parser.add_argument("-n", "--name", help="Schema name (default: filename)", default=None)
    args = parser.parse_args()

    input_path = Path(args.input)
    output_path = Path(args.output) if args.output else input_path.with_suffix(".schema.yml")
    schema_name = args.name or input_path.stem

    fields = extract_merge_fields(input_path)
    schema = generate_schema(fields, schema_name)
    save_schema(schema, output_path)

if __name__ == "__main__":
    main()
