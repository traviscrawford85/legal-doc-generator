import yaml
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from utils.style_applier import apply_styles

def load_yaml(path):
    with open(path, 'r') as f:
        return yaml.safe_load(f)

def set_font(run, style):
    run.font.name = style['font_name']
    run.font.size = Pt(style['font_size'])
    r = run._element
    r.rPr.rFonts.set(qn('w:eastAsia'), style['font_name'])

def build_lor_doc(context, style):
    doc = Document()
    apply_styles(doc, style)

    # VIA FAX LINE
    para = doc.add_paragraph()
    run = para.add_run(f"Via Fax: {context['Insurance']['AdjusterName']['EmailAddress']}")
    run.bold = True
    run.underline = True
    set_font(run, style)

    # INSURANCE ADDRESS
    for line in context['Insurance']['InsuranceAddress'].splitlines():
        p = doc.add_paragraph()
        run = p.add_run(line)
        set_font(run, style)

    doc.add_paragraph()  # spacer

    # Salutation
    para = doc.add_paragraph()
    run = para.add_run(f"Dear {context['Insurance']['AdjusterName']['FullName']},")
    set_font(run, style)

    # Body
    body = (
        f"Please be advised that our firm represents {context['Matter']['Client']['Name']} "
        f"in relation to injuries sustained in the motor vehicle accident on "
        f"{context['Matter']['Custom']['DateOfLoss']}."
    )
    para = doc.add_paragraph()
    run = para.add_run(body)
    set_font(run, style)

    insert_re_table(doc, context, style)
    doc.add_paragraph()  # Spacer
    insert_signature_block(doc, context, style)

    return doc

# Re: Table
def insert_re_table(doc, context, style):
    fields = [
        ("Claimant:", context["Matter"]["Client"]["Name"]),
        ("Date of Incident:", context["Matter"]["Custom"]["DateOfLoss"]),
        ("Claim No:", context["Matter"]["Custom"]["ClaimNumber"]),
        ("Policy No:", context["Matter"]["Custom"]["PolicyNumber"]),
    ]
    table = doc.add_table(rows=len(fields), cols=2)
    table.style = 'Table Grid'

    for i, (label, value) in enumerate(fields):
        cell_label = table.cell(i, 0)
        cell_value = table.cell(i, 1)

        p1 = cell_label.paragraphs[0]
        r1 = p1.add_run(label)
        r1.bold = True
        r1.font.name = style['font_name']
        r1.font.size = Pt(style['font_size'])
        r1._element.rPr.rFonts.set(qn('w:eastAsia'), style['font_name'])

        p2 = cell_value.paragraphs[0]
        r2 = p2.add_run(value)
        r2.font.name = style['font_name']
        r2.font.size = Pt(style['font_size'])
        r2._element.rPr.rFonts.set(qn('w:eastAsia'), style['font_name'])

def insert_signature_block(doc, context, style):
    doc.add_paragraph("Sincerely,")

    # Insert signature image
    sig_path = f"signatures/{context['Firm']['Attorney'].replace(' ', '')}.png"
    try:
        para = doc.add_paragraph()
        run = para.add_run()
        run.add_picture(sig_path, width=Inches(2))
    except Exception:
        doc.add_paragraph("[Signature Image Missing]")

    # Attorney name
    para = doc.add_paragraph()
    run = para.add_run(context['Firm']['Attorney'])
    run.font.name = style['font_name']
    run.font.size = Pt(style['font_size'])
    run._element.rPr.rFonts.set(qn('w:eastAsia'), style['font_name'])

    # Initials line
    para = doc.add_paragraph()
    run = para.add_run(f"PN/{context['User']['Initials']}")
    run.font.name = style['font_name']
    run.font.size = Pt(style['font_size'])
    run._element.rPr.rFonts.set(qn('w:eastAsia'), style['font_name'])

def main():
    context = load_yaml("content/example_context.yml")
    style = load_yaml("styles/representation_standard.yml")
    doc = build_lor_doc(context, style)
    doc.save("output/pip_letter.docx")

if __name__ == "__main__":
    main()
