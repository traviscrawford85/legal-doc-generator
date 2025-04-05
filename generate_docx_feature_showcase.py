from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

doc = Document()

# Set up document section
section = doc.sections[0]
section.top_margin = Inches(1)
section.bottom_margin = Inches(1)

# --- HEADER ---
header = section.header
header_para = header.paragraphs[0]
header_para.text = "Firm Header - Confidential"
header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

# --- FOOTER ---
footer = section.footer
footer.paragraphs[0].text = "Page 1 – Firm Contact | www.firmlaw.com"

# --- PARAGRAPH EXAMPLES ---
p1 = doc.add_paragraph("This is a normal paragraph.")
p1.paragraph_format.space_after = Pt(10)

p2 = doc.add_paragraph()
run = p2.add_run("This run is bold and blue.")
run.bold = True
run.font.color.rgb = RGBColor(0x00, 0x00, 0xFF)
run.font.size = Pt(12)

p3 = doc.add_paragraph()
run = p3.add_run("Georgia, italic, 14pt.")
run.font.name = "Georgia"
run.font.size = Pt(14)
run.italic = True
r = run._element
r.rPr.rFonts.set(qn('w:eastAsia'), 'Georgia')

# --- TABLE EXAMPLE ---
table = doc.add_table(rows=3, cols=2)
table.style = "Table Grid"
table.cell(0, 0).text = "Field"
table.cell(0, 1).text = "Value"
table.cell(1, 0).text = "Claimant"
table.cell(1, 1).text = "Jane Doe"
table.cell(2, 0).text = "Date of Loss"
table.cell(2, 1).text = "April 5, 2025"

# --- IMAGE EXAMPLE ---
doc.add_paragraph("Signature Example:")
doc.add_paragraph().add_run().add_picture("branding/logo.png", width=Inches(1.5))

# --- PAGE BREAK ---
doc.add_page_break()
doc.add_paragraph("This is the second page after the page break.")

# --- SAVE OUTPUT ---
doc.save("output/docx_feature_showcase.docx")
