from docx.shared import Inches, Pt
from docx.oxml.ns import qn

def insert_signature_block(doc, context, style):
    doc.add_paragraph("Sincerely,")

    # Insert signature image
    sig_path = f"signatures/{context['Firm']['Attorney'].replace(' ', '')}.png"
    try:
        para = doc.add_paragraph()
        run = para.add_run()
        run.add_picture(sig_path, width=Inches(2))
    except Exception:
        doc.add_paragraph("[Signature Image Missing]")

    # Attorney name
    para = doc.add_paragraph()
    run = para.add_run(context['Firm']['Attorney'])
    run.font.name = style['font_name']
    run.font.size = Pt(style['font_size'])
    run._element.rPr.rFonts.set(qn('w:eastAsia'), style['font_name'])

    # Initials line
    para = doc.add_paragraph()
    run = para.add_run(f"PN/{context['User']['Initials']}")
    run.font.name = style['font_name']
    run.font.size = Pt(style['font_size'])
    run._element.rPr.rFonts.set(qn('w:eastAsia'), style['font_name'])