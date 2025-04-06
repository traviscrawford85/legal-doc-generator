from docx.shared import Pt
from docx.oxml.ns import qn

def insert_re_table(doc, context, style):
    fields = [
        ("Claimant:", context["Matter"]["Client"]["Name"]),
        ("Date of Incident:", context["Matter"]["Custom"]["DateOfLoss"]),
        ("Claim No:", context["Matter"]["Custom"]["ClaimNumber"]),
        ("Policy No:", context["Matter"]["Custom"]["PolicyNumber"]),
    ]
    
    table = doc.add_table(rows=len(fields), cols=2)
    table.style = 'Table Grid'

    for i, (label, value) in enumerate(fields):
        cell_label = table.cell(i, 0)
        cell_value = table.cell(i, 1)

        # Label cell (left)
        p1 = cell_label.paragraphs[0]
        r1 = p1.add_run(label)
        r1.bold = True
        r1.font.name = style['font_name']
        r1.font.size = Pt(style['font_size'])
        r1._element.rPr.rFonts.set(qn('w:eastAsia'), style['font_name'])

        # Value cell (right)
        p2 = cell_value.paragraphs[0]
        r2 = p2.add_run(value)
        r2.font.name = style['font_name']
        r2.font.size = Pt(style['font_size'])
        r2._element.rPr.rFonts.set(qn('w:eastAsia'), style['font_name'])
