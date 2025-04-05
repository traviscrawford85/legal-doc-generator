from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

def insert_letterhead(doc, style):
    section = doc.sections[0]

    # Header: Add logo
    header = section.header
    if not header.paragraphs:
        header_para = header.add_paragraph()
    else:
        header_para = header.paragraphs[0]

    logo_run = header_para.add_run()
    try:
        logo_run.add_picture("branding/logo.png", width=Inches(1.5))
    except FileNotFoundError:
        header_para.add_run("[Missing logo.png]")

    # Footer: Firm contact info
    footer = section.footer
    for p in footer.paragraphs:
        p.clear()

    firm_info = [
        "Ledyard Law Group",
        "123 Legal Ave, Baltimore, MD 21201",
        "Phone: (410) 807-8077 | Fax: (410) 555-0111",
        "www.ledyardlaw.com"
    ]

    for line in firm_info:
        para = footer.add_paragraph()
        run = para.add_run(line)
        run.font.name = style['font_name']
        run.font.size = Pt(10)
        para.alignment = WD_ALIGN_PARAGRAPH.CENTER
