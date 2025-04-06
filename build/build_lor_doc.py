import sys
import os
sys.path.append(os.path.abspath(os.path.join(os.path.dirname(__file__), '..')))
import yaml
from docx import Document
from docx.shared import Pt, Inches
from docx.oxml.ns import qn
from utils.style_applier import apply_styles
from utils.re_table import insert_re_table
from utils.signature_block import insert_signature_block
from utils.schema_validator import load_schema, validate_context_against_schema

def load_yaml(path):
    with open(path, "r") as f:
        return yaml.safe_load(f)

def set_font(run, style):
    run.font.name = style['font_name']
    run.font.size = Pt(style['font_size'])
    run._element.rPr.rFonts.set(qn('w:eastAsia'), style['font_name'])

def build_lor_doc(context, style):
    doc = Document()
    apply_styles(doc, style)

    # Transmission block
    para = doc.add_paragraph()
    run = para.add_run(f"Via Fax: {context['Insurance']['AdjusterName']['EmailAddress']}")
    run.bold = True
    run.underline = True
    set_font(run, style)

    # Insurance address block
    for line in context['Insurance']['InsuranceAddress'].splitlines():
        p = doc.add_paragraph()
        run = p.add_run(line)
        set_font(run, style)

    doc.add_paragraph()  # Spacer

    # Salutation
    para = doc.add_paragraph()
    run = para.add_run(f"Dear {context['Insurance']['AdjusterName']['FullName']},")
    set_font(run, style)

    # Body
    para = doc.add_paragraph()
    body = (
        f"Please be advised that our firm represents {context['Matter']['Client']['Name']} "
        f"in relation to injuries sustained in the motor vehicle accident on "
        f"{context['Matter']['Custom']['DateOfLoss']}."
    )
    run = para.add_run(body)
    set_font(run, style)

    insert_re_table(doc, context, style)
    doc.add_paragraph()  # Spacer
    insert_signature_block(doc, context, style)

    return doc

def main():
    context = load_yaml("content/example_context.yml")
    style = load_yaml("styles/representation_standard.yml")
    schema = load_yaml("schemas/pip_lor.schema.yml")

    errors = validate_context_against_schema(context, schema)
    if errors:
        print("❌ Schema validation failed:")
        for err in errors:
            print(" -", err)
        exit(1)

    doc = build_lor_doc(context, style)
    doc.save("output/pip_letter.docx")
    print("✅ Document generated: output/pip_letter.docx")

if __name__ == "__main__":
    main()
